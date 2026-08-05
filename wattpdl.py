import os
import re
import sys
import html
import time
import pathlib
import tempfile
import zipfile
import requests

from rich.console import Console
from rich.theme import Theme
from rich.panel import Panel
from rich.table import Table
from rich.align import Align
from rich.progress import (
    Progress,
    SpinnerColumn,
    BarColumn,
    TextColumn,
    TimeElapsedColumn,
    TimeRemainingColumn,
)
from rich.prompt import Prompt
from rich.rule import Rule
from rich import box

THEME = Theme({
    "primary":   "bold cyan",
    "accent":    "cyan",
    "success":   "bold green",
    "warning":   "bold yellow",
    "danger":    "bold red",
    "muted":     "dim",
    "value":     "bold white",
})

console = Console(theme=THEME)
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )
}

STORY_INFO_URL = (
    "https://www.wattpad.com/api/v3/stories/{story_id}"
    "?fields=id,title,user(name),numParts,parts(id,title)"
)
CHAPTER_TEXT_URL = "https://www.wattpad.com/apiv2/storytext?id={part_id}"

DELAY_SECONDS = 0.5   # jeda antar chapter (jangan dihapus)
MAX_RETRIES   = 3    

LOGO = r"""
[primary]██╗    ██╗ █████╗ ████████╗████████╗██████╗ ██████╗ ██╗     [/primary]
[primary]██║    ██║██╔══██╗╚══██╔══╝╚══██╔══╝██╔══██╗██╔══██╗██║     [/primary]
[primary]██║ █╗ ██║███████║   ██║      ██║   ██████╔╝██║  ██║██║     [/primary]
[primary]██║███╗██║██╔══██║   ██║      ██║   ██╔═══╝ ██║  ██║██║     [/primary]
[primary]╚███╔███╔╝██║  ██║   ██║      ██║   ██║     ██████╔╝███████╗[/primary]
[muted] ╚══╝╚══╝ ╚═╝  ╚═╝   ╚═╝      ╚═╝   ╚═╝     ╚═════╝ ╚══════╝[/muted]
"""

_step_counter = {"n": 0}
_CIRCLED = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩"]


def step_rule(title: str) -> None:
    """Tampilkan pembatas antar tahap dengan nomor otomatis."""
    idx = _step_counter["n"]
    num = _CIRCLED[idx] if idx < len(_CIRCLED) else str(idx + 1)
    _step_counter["n"] += 1
    console.print()
    console.print(Rule(f"[primary]{num}  {title}[/primary]", style="accent", align="left"))


def get_default_download_dir() -> pathlib.Path:
    """
    Kembalikan folder Downloads bawaan sistem operasi:
      - Windows : C:\\Users\\<user>\\Downloads
      - macOS   : /Users/<user>/Downloads
      - Linux   : /home/<user>/Downloads  (atau XDG_DOWNLOAD_DIR)
    """
    xdg = os.environ.get("XDG_DOWNLOAD_DIR", "").strip()
    if xdg and pathlib.Path(xdg).is_dir():
        return pathlib.Path(xdg)

    downloads = pathlib.Path.home() / "Downloads"
    downloads.mkdir(parents=True, exist_ok=True)
    return downloads


def extract_story_id(user_input: str) -> str:
    """Terima link penuh atau ID angka, kembalikan ID cerita saja."""
    user_input = user_input.strip()
    match = re.search(r"story/(\d+)", user_input)
    if match:
        return match.group(1)
    if user_input.isdigit():
        return user_input
    raise ValueError(
        "Link atau ID cerita tidak dikenali.\n"
        "Contoh link : https://www.wattpad.com/story/398440633-judul-cerita\n"
        "Contoh ID   : 398440633"
    )


def get_story_info(story_id: str) -> tuple[str, str, list]:
    """Ambil metadata cerita: judul, penulis, dan daftar chapter."""
    url  = STORY_INFO_URL.format(story_id=story_id)
    resp = requests.get(url, headers=HEADERS, timeout=20)
    resp.raise_for_status()
    data   = resp.json()
    title  = data.get("title", f"story_{story_id}")
    author = data.get("user", {}).get("name", "unknown")
    parts  = data.get("parts", [])
    return title, author, parts


def get_chapter_html(part_id: int, retries: int = MAX_RETRIES) -> str:
    """Unduh HTML teks chapter dengan retry otomatis."""
    url = CHAPTER_TEXT_URL.format(part_id=part_id)
    for attempt in range(1, retries + 1):
        try:
            resp = requests.get(url, headers=HEADERS, timeout=20)
            resp.raise_for_status()
            return resp.text
        except requests.RequestException as e:
            if attempt == retries:
                raise
            wait = attempt * 2
            console.print(
                f"    [warning]⚠  Gagal (percobaan {attempt}/{retries}), "
                f"coba lagi dalam {wait}s… ({e})[/warning]"
            )
            time.sleep(wait)
    return ""   


def html_to_text(raw_html: str) -> str:
    """Konversi HTML Wattpad ke teks bersih."""
    text = re.sub(r"</p>",       "\n\n", raw_html, flags=re.IGNORECASE)
    text = re.sub(r"<br\s*/?>",  "\n",   text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>",   "",     text)
    text = html.unescape(text)
    text = re.sub(r"\n{3,}",    "\n\n", text)
    return text.strip()


def safe_filename(name: str) -> str:
    """Buat nama file yang aman untuk semua sistem operasi."""
    name = re.sub(r'[<>:"/\\|?*]', "", name)
    name = re.sub(r"[^\w\s\-]", "", name).strip()
    name = re.sub(r"\s+", "_", name)
    return name or "cerita_wattpad"


def format_duration(seconds: float) -> str:
    """Ubah detik jadi format menit:detik yang enak dibaca."""
    seconds = int(seconds)
    minutes, secs = divmod(seconds, 60)
    if minutes:
        return f"{minutes}m {secs}d"
    return f"{secs}d"


def check_docx_available() -> None:
    """Pastikan library python-docx terinstall sebelum dipakai."""
    try:
        import docx  
    except ImportError:
        console.print("\n[danger]❌  Library 'python-docx' belum terinstall.[/danger]")
        console.print("    [muted]Jalankan:[/muted] [accent]pip install python-docx[/accent]")
        sys.exit(1)


def show_chapter_list(parts: list) -> None:
    """Tampilkan tabel nomor + judul semua chapter."""
    table = Table(box=box.SIMPLE_HEAD, padding=(0, 1), expand=False)
    table.add_column("No", style="muted", justify="right", width=5)
    table.add_column("Judul Chapter", style="value")
    for i, part in enumerate(parts, start=1):
        table.add_row(str(i), part.get("title", f"Chapter {i}"))
    console.print(table)


def parse_chapter_selection(raw: str, total: int) -> list:
    """Terima input seperti '1,3,5-8' dan kembalikan list nomor chapter unik & urut."""
    raw = raw.replace(" ", "")
    if not raw:
        raise ValueError("Input tidak boleh kosong.")

    selected = set()
    for chunk in raw.split(","):
        if not chunk:
            continue
        if "-" in chunk:
            start_s, _, end_s = chunk.partition("-")
            if not start_s.isdigit() or not end_s.isdigit():
                raise ValueError(f"Format rentang tidak valid: '{chunk}'")
            start, end = int(start_s), int(end_s)
            if start > end:
                start, end = end, start
            selected.update(n for n in range(start, end + 1) if 1 <= n <= total)
        else:
            if not chunk.isdigit():
                raise ValueError(f"'{chunk}' bukan angka yang valid.")
            n = int(chunk)
            if 1 <= n <= total:
                selected.add(n)

    if not selected:
        raise ValueError(f"Tidak ada nomor chapter valid (rentang 1-{total}).")
    return sorted(selected)


def select_multiple_chapters(parts: list) -> list:
    """Tampilkan daftar chapter, minta user memilih beberapa. Kembalikan list (nomor, part)."""
    step_rule("Pilih Chapter")
    show_chapter_list(parts)
    while True:
        raw = Prompt.ask(
            "\n[value]Masukkan nomor chapter[/value] [muted](contoh: 1,3,5-8)[/muted]"
        )
        try:
            numbers = parse_chapter_selection(raw, len(parts))
            break
        except ValueError as e:
            console.print(f"[danger]❌  {e}[/danger]")

    console.print(f"[success]✔ {len(numbers)} chapter dipilih.[/success]")
    return [(n, parts[n - 1]) for n in numbers]


def select_single_chapter(parts: list) -> tuple:
    """Tampilkan daftar chapter, minta user memilih satu. Kembalikan (nomor, part)."""
    step_rule("Pilih Chapter")
    show_chapter_list(parts)
    while True:
        raw = Prompt.ask("\n[value]Masukkan nomor chapter[/value]").strip()
        if raw.isdigit() and 1 <= int(raw) <= len(parts):
            n = int(raw)
            break
        console.print(f"[danger]❌  Masukkan angka antara 1-{len(parts)}.[/danger]")

    part = parts[n - 1]
    console.print(f"[success]✔ Dipilih:[/success] Chapter {n} — {part.get('title', '')}")
    return n, part


def download_chapters(indexed_parts: list) -> tuple:
    """
    Unduh sekumpulan chapter dengan progress bar.
    indexed_parts: list of (nomor_chapter, part_dict)
    Return: (results, failed_chapters)
      results = list of (nomor_chapter, judul_chapter, teks)
    """
    results = []
    failed  = []

    progress_columns = [
        SpinnerColumn(style="accent", spinner_name="dots"),
        TextColumn("[value]{task.description}[/value]"),
        BarColumn(bar_width=32, complete_style="cyan", finished_style="success", style="grey23"),
        TextColumn("[progress.percentage]{task.percentage:>5.1f}%"),
        TextColumn("[muted]{task.completed}/{task.total}[/muted]"),
        TextColumn("[muted]•[/muted]"),
        TimeElapsedColumn(),
        TextColumn("[muted]sisa[/muted]"),
        TimeRemainingColumn(),
    ]

    with Progress(*progress_columns, console=console) as progress:
        task = progress.add_task("Menyiapkan…", total=len(indexed_parts))

        for idx, part in indexed_parts:
            part_id       = part["id"]
            chapter_title = part.get("title", f"Chapter {idx}")

            short_title = (chapter_title[:26] + "…") if len(chapter_title) > 26 else chapter_title
            progress.update(task, description=f"{short_title:<28}")

            try:
                html = get_chapter_html(part_id)
                text = html_to_text(html)
            except Exception as e:
                console.print(f"\n    [danger]⚠  Chapter [{idx}] gagal:[/danger] {e}")
                text = "[GAGAL DIUNDUH — coba jalankan ulang]"
                failed.append(f"Chapter {idx}: {chapter_title}")

            results.append((idx, chapter_title, text))
            progress.advance(task)
            time.sleep(DELAY_SECONDS)

    return results, failed


def write_combined_txt(path: pathlib.Path, title: str, author: str, story_id: str, results: list) -> None:
    """Tulis semua chapter yang diberikan ke satu file .txt."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"{title}\n")
        f.write(f"oleh {author}\n")
        f.write(f"Sumber : https://www.wattpad.com/story/{story_id}\n")
        f.write(f"\n{'=' * 50}\n\n")
        for _, chapter_title, text in results:
            f.write(f"\n\n{'#' * 5} {chapter_title} {'#' * 5}\n\n")
            f.write(text)
            f.write("\n")


def write_separate_zip(path: pathlib.Path, title: str, author: str, story_id: str, results: list) -> None:
    """Tulis tiap chapter sebagai file .txt terpisah, dikemas dalam satu .zip."""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        info = (
            f"{title}\n"
            f"oleh {author}\n"
            f"Sumber : https://www.wattpad.com/story/{story_id}\n"
        )
        zf.writestr("000_info.txt", info)
        for idx, chapter_title, text in results:
            fname = f"{idx:03d}_{safe_filename(chapter_title)}.txt"
            zf.writestr(fname, f"{chapter_title}\n\n{text}\n")


def _write_paragraphs(doc, text: str) -> None:
    """Tulis teks chapter ke dokumen Word, menjaga jeda antar paragraf."""
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        lines = para.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            p.add_run(line)


def write_combined_docx(path: pathlib.Path, title: str, author: str, story_id: str, results: list) -> None:
    """Tulis semua chapter yang diberikan ke satu file .docx."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"oleh {author}")
    doc.add_paragraph(f"Sumber : https://www.wattpad.com/story/{story_id}")

    for i, (_, chapter_title, text) in enumerate(results):
        doc.add_page_break()
        doc.add_heading(chapter_title, level=1)
        _write_paragraphs(doc, text)

    doc.save(str(path))


def write_separate_docx_zip(path: pathlib.Path, title: str, author: str, story_id: str, results: list) -> None:
    """Tulis tiap chapter sebagai file .docx terpisah, dikemas dalam satu .zip."""
    from docx import Document

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = pathlib.Path(tmpdir)

        info_doc = Document()
        info_doc.add_heading(title, level=0)
        info_doc.add_paragraph(f"oleh {author}")
        info_doc.add_paragraph(f"Sumber : https://www.wattpad.com/story/{story_id}")
        info_file = tmp_path / "000_info.docx"
        info_doc.save(str(info_file))

        chapter_files = [info_file]
        for idx, chapter_title, text in results:
            d = Document()
            d.add_heading(chapter_title, level=1)
            _write_paragraphs(d, text)
            fname = tmp_path / f"{idx:03d}_{safe_filename(chapter_title)}.docx"
            d.save(str(fname))
            chapter_files.append(fname)

        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            for f in chapter_files:
                zf.write(f, arcname=f.name)


def main():
    console.print()
    console.print(Align.center(LOGO))
    console.print(
        Align.center(
            "[value]📖 Wattpad Story Downloader[/value]  [muted]—  simpan cerita jadi file .txt/.docx offline[/muted]"
        )
    )

    step_rule("Masukkan Cerita")
    raw = Prompt.ask("[value]Link atau ID cerita Wattpad[/value]")
    try:
        story_id = extract_story_id(raw)
    except ValueError as e:
        console.print(f"\n[danger]❌  {e}[/danger]")
        sys.exit(1)

    with console.status(f"[accent]Mengambil info cerita (ID: {story_id})…[/accent]", spinner="dots"):
        try:
            title, author, parts = get_story_info(story_id)
        except requests.HTTPError as e:
            console.print(f"[danger]❌  Gagal mengakses API Wattpad:[/danger] {e}")
            console.print("    [muted]Pastikan ID/link benar dan cerita tidak di-private.[/muted]")
            sys.exit(1)

    if not parts:
        console.print("[danger]❌  Tidak ada chapter ditemukan. Cek lagi ID/link ceritanya.[/danger]")
        sys.exit(1)

    step_rule("Info Cerita")
    info_table = Table(show_header=False, box=box.SIMPLE, padding=(0, 1), expand=False)
    info_table.add_column(style="muted")
    info_table.add_column(style="value")
    info_table.add_row("📌 Judul", title)
    info_table.add_row("✍️  Penulis", author)
    info_table.add_row("📚 Jumlah chapter", str(len(parts)))
    console.print(Panel(info_table, border_style="success", box=box.ROUNDED, padding=(1, 2)))

    step_rule("Pilih Mode Unduh")
    mode_table = Table(show_header=False, box=box.SIMPLE, padding=(0, 1), expand=False)
    mode_table.add_column(style="accent", justify="right")
    mode_table.add_column(style="value")
    mode_table.add_row("1", "Semua chapter → 1 file gabungan")
    mode_table.add_row("2", "Semua chapter → file terpisah per chapter, dikemas .zip")
    mode_table.add_row("3", "Pilih beberapa chapter → 1 file gabungan")
    mode_table.add_row("4", "Pilih 1 chapter saja")
    console.print(mode_table)

    mode = Prompt.ask(
        "\n[value]Pilih mode[/value]",
        choices=["1", "2", "3", "4"],
        default="1",
    )

    format_table = Table(show_header=False, box=box.SIMPLE, padding=(0, 1), expand=False)
    format_table.add_column(style="accent", justify="right")
    format_table.add_column(style="value")
    format_table.add_row("1", "Teks polos (.txt)")
    format_table.add_row("2", "Dokumen Word (.docx)")
    console.print()
    console.print(format_table)

    file_format = Prompt.ask(
        "\n[value]Pilih format file[/value]",
        choices=["1", "2"],
        default="1",
    )
    if file_format == "2":
        check_docx_available()

    if mode == "3":
        indexed_parts = select_multiple_chapters(parts)
    elif mode == "4":
        chosen = select_single_chapter(parts)
        indexed_parts = [chosen]
    else:
        indexed_parts = list(enumerate(parts, start=1))

    step_rule("Folder Penyimpanan")
    default_dir = get_default_download_dir()
    console.print(f"[muted]Folder default:[/muted] [accent]{default_dir}[/accent]")
    custom = Prompt.ask(
        "[muted]Tekan Enter untuk pakai folder itu, atau ketik path lain[/muted]",
        default="",
        show_default=False,
    ).strip()

    if custom:
        save_dir = pathlib.Path(custom).expanduser().resolve()
        try:
            save_dir.mkdir(parents=True, exist_ok=True)
        except OSError as e:
            console.print(f"[danger]❌  Folder tidak bisa dibuat:[/danger] {e}")
            console.print(f"    [muted]Menggunakan folder default: {default_dir}[/muted]")
            save_dir = default_dir
    else:
        save_dir = default_dir

    ext = "docx" if file_format == "2" else "txt"
    base_name = safe_filename(title)

    if mode == "1":
        full_path = save_dir / f"{base_name}.{ext}"
    elif mode == "2":
        full_path = save_dir / f"{base_name}.zip"
    elif mode == "3":
        full_path = save_dir / f"{base_name}_pilihan.{ext}"
    else:  
        chap_no, chap_part = indexed_parts[0]
        chap_title = safe_filename(chap_part.get("title", f"Chapter {chap_no}"))
        full_path = save_dir / f"{base_name}_Ch{chap_no:03d}_{chap_title}.{ext}"

    console.print(f"\n[value]💾 File akan disimpan di:[/value] [accent]{full_path}[/accent]")

    step_rule("Mengunduh Chapter")
    start_time = time.time()
    results, failed_chapters = download_chapters(indexed_parts)
    elapsed = format_duration(time.time() - start_time)

    if mode == "2":
        if file_format == "2":
            write_separate_docx_zip(full_path, title, author, story_id, results)
        else:
            write_separate_zip(full_path, title, author, story_id, results)
    else:
        if file_format == "2":
            write_combined_docx(full_path, title, author, story_id, results)
        else:
            write_combined_txt(full_path, title, author, story_id, results)

    step_rule("Ringkasan")

    summary = Table(show_header=False, box=box.SIMPLE, padding=(0, 1), expand=False)
    summary.add_column(style="muted")
    summary.add_column(style="value")
    summary.add_row("📁 File tersimpan di", f"[accent]{full_path}[/accent]")
    summary.add_row("✅ Berhasil", f"[success]{len(results) - len(failed_chapters)}/{len(results)}[/success]")
    if failed_chapters:
        summary.add_row("⚠️  Gagal", f"[danger]{len(failed_chapters)}[/danger]")
    summary.add_row("⏱️  Waktu total", elapsed)

    if failed_chapters:
        fail_list = "\n".join(f"[danger]• {ch}[/danger]" for ch in failed_chapters)
        body = Table.grid(padding=(0, 0))
        body.add_row(summary)
        body.add_row("")
        body.add_row(f"[muted]Chapter yang gagal:[/muted]\n{fail_list}")
        console.print(Panel(
            body,
            title="[warning]⚠️  Selesai dengan beberapa kegagalan[/warning]",
            border_style="warning",
            box=box.ROUNDED,
            padding=(1, 2),
        ))
        console.print("[muted]Jalankan ulang script untuk mencoba mengunduh ulang chapter yang gagal.[/muted]\n")
    else:
        console.print(Panel(
            summary,
            title="[success]🎉 Berhasil!  Semua chapter tersimpan[/success]",
            border_style="success",
            box=box.DOUBLE,
            padding=(1, 2),
        ))
        console.print()

if __name__ == "__main__":
    main()
