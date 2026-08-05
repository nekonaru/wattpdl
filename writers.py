"""
Modul konversi teks dan penulisan file output.
Berisi semua fungsi yang mengubah teks chapter mentah menjadi file
.txt / .docx / .zip di disk. Tidak ada logika jaringan atau tampilan CLI di sini.
"""
import html
import pathlib
import re
import sys
import tempfile
import zipfile


def html_to_text(raw_html: str) -> str:
    """Konversi HTML Wattpad ke teks bersih."""
    text = re.sub(r"</p>",       "\n\n", raw_html, flags=re.IGNORECASE)
    text = re.sub(r"<br\s*/?>",  "\n",   text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>",   "",     text)
    text = html.unescape(text)
    text = re.sub(r"\n{3,}",    "\n\n", text)
    return text.strip()


def safe_filename(name: str) -> str:
    """Buat nama file yang aman untuk semua sistem operasi."""
    name = re.sub(r'[<>:"/\\|?*]', "", name)
    name = re.sub(r"[^\w\s\-]", "", name).strip()
    name = re.sub(r"\s+", "_", name)
    return name or "cerita_wattpad"


def check_docx_available(console=None) -> None:
    """Pastikan library python-docx terinstall sebelum dipakai."""
    try:
        import docx  # noqa: F401
    except ImportError:
        if console:
            console.print("\n[danger]❌  Library 'python-docx' belum terinstall.[/danger]")
            console.print("    [muted]Jalankan:[/muted] [accent]pip install python-docx[/accent]")
        else:
            print("Library 'python-docx' belum terinstall. Jalankan: pip install python-docx")
        sys.exit(1)


def write_combined_txt(path: pathlib.Path, title: str, author: str, story_id: str, results: list) -> None:
    """Tulis semua chapter yang diberikan ke satu file .txt."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"{title}\n")
        f.write(f"oleh {author}\n")
        f.write(f"Sumber : https://www.wattpad.com/story/{story_id}\n")
        f.write(f"\n{'=' * 50}\n\n")
        for _, chapter_title, text in results:
            f.write(f"\n\n{'#' * 5} {chapter_title} {'#' * 5}\n\n")
            f.write(text)
            f.write("\n")


def write_separate_zip(path: pathlib.Path, title: str, author: str, story_id: str, results: list) -> None:
    """Tulis tiap chapter sebagai file .txt terpisah, dikemas dalam satu .zip."""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        info = (
            f"{title}\n"
            f"oleh {author}\n"
            f"Sumber : https://www.wattpad.com/story/{story_id}\n"
        )
        zf.writestr("000_info.txt", info)
        for idx, chapter_title, text in results:
            fname = f"{idx:03d}_{safe_filename(chapter_title)}.txt"
            zf.writestr(fname, f"{chapter_title}\n\n{text}\n")


def _write_paragraphs(doc, text: str) -> None:
    """Tulis teks chapter ke dokumen Word, menjaga jeda antar paragraf."""
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        lines = para.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            p.add_run(line)


def write_combined_docx(path: pathlib.Path, title: str, author: str, story_id: str, results: list) -> None:
    """Tulis semua chapter yang diberikan ke satu file .docx."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"oleh {author}")
    doc.add_paragraph(f"Sumber : https://www.wattpad.com/story/{story_id}")

    for _, chapter_title, text in results:
        doc.add_page_break()
        doc.add_heading(chapter_title, level=1)
        _write_paragraphs(doc, text)

    doc.save(str(path))


def write_separate_docx_zip(path: pathlib.Path, title: str, author: str, story_id: str, results: list) -> None:
    """Tulis tiap chapter sebagai file .docx terpisah, dikemas dalam satu .zip."""
    from docx import Document

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = pathlib.Path(tmpdir)

        info_doc = Document()
        info_doc.add_heading(title, level=0)
        info_doc.add_paragraph(f"oleh {author}")
        info_doc.add_paragraph(f"Sumber : https://www.wattpad.com/story/{story_id}")
        info_file = tmp_path / "000_info.docx"
        info_doc.save(str(info_file))

        chapter_files = [info_file]
        for idx, chapter_title, text in results:
            d = Document()
            d.add_heading(chapter_title, level=1)
            _write_paragraphs(d, text)
            fname = tmp_path / f"{idx:03d}_{safe_filename(chapter_title)}.docx"
            d.save(str(fname))
            chapter_files.append(fname)

        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            for f in chapter_files:
                zf.write(f, arcname=f.name)
