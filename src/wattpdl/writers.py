"""
Modul konversi teks dan penulisan file output.
Berisi semua fungsi yang mengubah teks chapter mentah menjadi file
.txt / .docx / .zip di disk. Tidak ada logika jaringan atau tampilan CLI di sini.
"""
import html
import pathlib
import re
import sys
import tempfile
import zipfile


def html_to_text(raw_html: str) -> str:
    """Konversi HTML Wattpad ke teks bersih."""
    text = re.sub(r"</p>",       "\n\n", raw_html, flags=re.IGNORECASE)
    text = re.sub(r"<br\s*/?>",  "\n",   text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>",   "",     text)
    text = html.unescape(text)
    text = re.sub(r"\n{3,}",    "\n\n", text)
    return text.strip()


_WINDOWS_RESERVED_NAMES = {
    "CON", "PRN", "AUX", "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}


def safe_filename(name: str, max_length: int = 100) -> str:
    """Buat nama file yang aman untuk semua sistem operasi.

    Judul cerita/chapter Wattpad kadang sangat panjang (umum untuk judul
    bergaya panjang/"clickbait"). Dipotong ke `max_length` supaya nama file
    akhir (yang di beberapa tempat menggabungkan judul cerita + judul chapter,
    lihat mode 4 di app.py) tidak melebihi batas panjang nama file sistem
    operasi (mis. 255 byte di Linux/macOS) dan menyebabkan OSError
    "File name too long" yang tidak tertangani saat menyimpan.

    Juga menghindari nama device reserved Windows (CON, AUX, PRN, NUL,
    COM1-9, LPT1-9) — kalau judul chapter kebetulan cuma satu kata seperti
    itu (jarang, tapi mungkin), menyimpan file dengan nama itu akan selalu
    gagal di Windows apa pun ekstensinya.
    """
    name = re.sub(r'[<>:"/\\|?*]', "", name)
    name = re.sub(r"[^\w\s\-]", "", name).strip()
    name = re.sub(r"\s+", "_", name)
    name = name[:max_length].rstrip("_")
    if not name:
        return "cerita_wattpad"
    if name.upper() in _WINDOWS_RESERVED_NAMES:
        name = f"_{name}"
    return name


def check_docx_available(console=None) -> None:
    """Pastikan library python-docx terinstall sebelum dipakai."""
    try:
        import docx  # noqa: F401
    except ImportError:
        if console:
            console.print("\n[danger]❌  Library 'python-docx' belum terinstall.[/danger]")
            console.print("    [muted]Jalankan:[/muted] [accent]pip install python-docx[/accent]")
        else:
            print("Library 'python-docx' belum terinstall. Jalankan: pip install python-docx")
        sys.exit(1)


def check_epub_available(console=None) -> None:
    """Pastikan library EbookLib terinstall sebelum dipakai."""
    try:
        import ebooklib  # noqa: F401
    except ImportError:
        if console:
            console.print("\n[danger]❌  Library 'EbookLib' belum terinstall.[/danger]")
            console.print("    [muted]Jalankan:[/muted] [accent]pip install EbookLib[/accent]")
        else:
            print("Library 'EbookLib' belum terinstall. Jalankan: pip install EbookLib")
        sys.exit(1)


def write_combined_txt(path: pathlib.Path, title: str, author: str, story_id: str, results: list) -> None:
    """Tulis semua chapter yang diberikan ke satu file .txt."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"{title}\n")
        f.write(f"oleh {author}\n")
        f.write(f"Sumber : https://www.wattpad.com/story/{story_id}\n")
        f.write(f"\n{'=' * 50}\n\n")
        for _, chapter_title, text in results:
            f.write(f"\n\n{'#' * 5} {chapter_title} {'#' * 5}\n\n")
            f.write(text)
            f.write("\n")


def write_combined_md(path: pathlib.Path, title: str, author: str, story_id: str, results: list,
                       meta: dict = None) -> None:
    """Tulis semua chapter yang diberikan ke satu file Markdown (.md)."""
    meta = meta or {}
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"# {title}\n\n")
        f.write(f"**oleh** {author}\n\n")
        if meta.get("tags"):
            f.write(f"**Genre**: {', '.join(meta['tags'])}\n\n")
        if meta.get("description"):
            f.write(f"> {meta['description']}\n\n")
        f.write(f"Sumber: <https://www.wattpad.com/story/{story_id}>\n\n")
        f.write("---\n")
        for _, chapter_title, text in results:
            f.write(f"\n## {chapter_title}\n\n")
            for para in text.split("\n\n"):
                para = para.strip()
                if para:
                    f.write(f"{para}\n\n")


def write_separate_md_zip(path: pathlib.Path, title: str, author: str, story_id: str, results: list,
                           meta: dict = None) -> None:
    """Tulis tiap chapter sebagai file .md terpisah, dikemas dalam satu .zip."""
    meta = meta or {}
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        info_lines = [f"# {title}", "", f"**oleh** {author}"]
        if meta.get("tags"):
            info_lines.append(f"**Genre**: {', '.join(meta['tags'])}")
        if meta.get("description"):
            info_lines.append(f"\n> {meta['description']}")
        info_lines.append(f"\nSumber: <https://www.wattpad.com/story/{story_id}>")
        zf.writestr("000_info.md", "\n".join(info_lines) + "\n")
        for idx, chapter_title, text in results:
            fname = f"{idx:03d}_{safe_filename(chapter_title)}.md"
            body = "\n\n".join(p.strip() for p in text.split("\n\n") if p.strip())
            zf.writestr(fname, f"## {chapter_title}\n\n{body}\n")


def write_separate_zip(path: pathlib.Path, title: str, author: str, story_id: str, results: list) -> None:
    """Tulis tiap chapter sebagai file .txt terpisah, dikemas dalam satu .zip."""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        info = (
            f"{title}\n"
            f"oleh {author}\n"
            f"Sumber : https://www.wattpad.com/story/{story_id}\n"
        )
        zf.writestr("000_info.txt", info)
        for idx, chapter_title, text in results:
            fname = f"{idx:03d}_{safe_filename(chapter_title)}.txt"
            zf.writestr(fname, f"{chapter_title}\n\n{text}\n")


def _write_paragraphs(doc, text: str) -> None:
    """Tulis teks chapter ke dokumen Word, menjaga jeda antar paragraf."""
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        lines = para.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            p.add_run(line)


def _insert_docx_cover(doc, cover_bytes: bytes) -> None:
    """Sisipkan gambar sampul di awal dokumen Word. Gagal decode/insert tidak boleh
    menggagalkan seluruh proses penulisan file — cover cuma pemanis."""
    import io

    from docx.shared import Inches

    try:
        doc.add_picture(io.BytesIO(cover_bytes), width=Inches(4))
    except Exception:
        pass


def _insert_docx_images(doc, images: list) -> None:
    """Sisipkan sekumpulan gambar inline (dari dalam teks chapter, bukan cover)
    ke akhir bagian chapter yang sedang ditulis. Sama seperti cover: gagal
    decode 1 gambar tidak boleh menggagalkan gambar lain / seluruh proses,
    jadi tiap gambar dicoba satu-satu dan yang gagal dilewati saja."""
    import io

    from docx.shared import Inches

    for img_bytes in images or []:
        try:
            doc.add_picture(io.BytesIO(img_bytes), width=Inches(4))
        except Exception:
            continue


def write_combined_docx(path: pathlib.Path, title: str, author: str, story_id: str, results: list,
                         meta: dict = None, cover_bytes: bytes = None, images_by_idx: dict = None) -> None:
    """Tulis semua chapter yang diberikan ke satu file .docx.
    images_by_idx: opsional {nomor_chapter: [bytes_gambar, ...]} — gambar inline
    dari dalam teks chapter (bukan cover), disisipkan di akhir tiap chapter.
    """
    from docx import Document

    meta = meta or {}
    images_by_idx = images_by_idx or {}
    doc = Document()
    if cover_bytes:
        _insert_docx_cover(doc, cover_bytes)
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"oleh {author}")
    if meta.get("tags"):
        doc.add_paragraph(f"Genre: {', '.join(meta['tags'])}")
    if meta.get("description"):
        doc.add_paragraph(meta["description"])
    doc.add_paragraph(f"Sumber : https://www.wattpad.com/story/{story_id}")

    for idx, chapter_title, text in results:
        doc.add_page_break()
        doc.add_heading(chapter_title, level=1)
        _write_paragraphs(doc, text)
        _insert_docx_images(doc, images_by_idx.get(idx))

    doc.save(str(path))


def write_separate_docx_zip(path: pathlib.Path, title: str, author: str, story_id: str, results: list,
                             meta: dict = None, cover_bytes: bytes = None, images_by_idx: dict = None) -> None:
    """Tulis tiap chapter sebagai file .docx terpisah, dikemas dalam satu .zip."""
    from docx import Document

    meta = meta or {}
    images_by_idx = images_by_idx or {}
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = pathlib.Path(tmpdir)

        info_doc = Document()
        if cover_bytes:
            _insert_docx_cover(info_doc, cover_bytes)
        info_doc.add_heading(title, level=0)
        info_doc.add_paragraph(f"oleh {author}")
        if meta.get("tags"):
            info_doc.add_paragraph(f"Genre: {', '.join(meta['tags'])}")
        if meta.get("description"):
            info_doc.add_paragraph(meta["description"])
        info_doc.add_paragraph(f"Sumber : https://www.wattpad.com/story/{story_id}")
        info_file = tmp_path / "000_info.docx"
        info_doc.save(str(info_file))

        chapter_files = [info_file]
        for idx, chapter_title, text in results:
            d = Document()
            d.add_heading(chapter_title, level=1)
            _write_paragraphs(d, text)
            _insert_docx_images(d, images_by_idx.get(idx))
            fname = tmp_path / f"{idx:03d}_{safe_filename(chapter_title)}.docx"
            d.save(str(fname))
            chapter_files.append(fname)

        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            for f in chapter_files:
                zf.write(f, arcname=f.name)


def _text_to_epub_html(text: str) -> str:
    """Ubah teks chapter (paragraf dipisah \\n\\n, baris dipisah \\n) jadi HTML aman untuk EPUB."""
    paragraphs = []
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        lines = [html.escape(line) for line in para.split("\n")]
        paragraphs.append(f"<p>{'<br/>'.join(lines)}</p>")
    return "\n".join(paragraphs)


def _build_epub_book(title: str, author: str, story_id: str, identifier_suffix: str = None,
                      meta: dict = None, language: str = "id"):
    """
    Buat objek EpubBook kosong dengan metadata standar wattpdl.
    identifier_suffix: opsional, dipakai untuk membuat identifier EPUB unik per file
    (mis. per chapter) — tanpa ini, semua EPUB dari cerita yang sama akan punya
    identifier persis sama, yang melanggar spec EPUB dan bisa bikin sebagian
    e-reader (Calibre, Kobo) bentrok/menimpa entri satu sama lain.
    meta: dict opsional {"description", "tags", "create_date"} — kalau diisi,
    ditambahkan sebagai metadata Dublin Core standar (description, subject, date)
    supaya kelihatan di info buku pada e-reader yang mendukungnya.
    language: kode bahasa ISO 639-1 utk metadata EPUB (mis. "id", "en"). Default
    "id" karena wattpdl ditujukan utk pembaca Wattpad Indonesia, tapi Wattpad
    juga menghosting banyak cerita berbahasa lain — kalau ceritanya bahasa
    Inggris, isi eksplisit lewat --lang supaya metadata EPUB-nya akurat
    (bukan cuma tampilan/estetika: sebagian e-reader memakai field ini utk
    text-to-speech dan sortir rak buku per bahasa).
    """
    from ebooklib import epub

    identifier = f"wattpdl-{story_id}" if identifier_suffix is None else f"wattpdl-{story_id}-{identifier_suffix}"
    book = epub.EpubBook()
    book.set_identifier(identifier)
    book.set_title(title)
    book.set_language(language or "id")
    book.add_author(author)

    meta = meta or {}
    if meta.get("description"):
        book.add_metadata("DC", "description", meta["description"])
    for tag in meta.get("tags") or []:
        book.add_metadata("DC", "subject", tag)
    if meta.get("create_date"):
        book.add_metadata("DC", "date", str(meta["create_date"]))

    return book, epub


def _set_epub_cover(book, epub, cover_bytes: bytes) -> None:
    """Sisipkan gambar sampul ke EpubBook. Gagal insert tidak boleh menggagalkan
    seluruh proses penulisan file — cover cuma pemanis."""
    if not cover_bytes:
        return
    try:
        book.set_cover("cover.jpg", cover_bytes)
    except Exception:
        pass


def _add_epub_inline_images(book, epub, idx: int, images: list) -> str:
    """Daftarkan sekumpulan gambar inline (dari dalam teks chapter) sebagai
    EpubImage ke buku, dan kembalikan potongan HTML <img> untuk disisipkan di
    akhir konten chapter. Gagal 1 gambar tidak menggagalkan gambar lain."""
    if not images:
        return ""
    tags = []
    for n, img_bytes in enumerate(images, start=1):
        try:
            file_name = f"images/ch{idx:03d}_img{n:02d}.jpg"
            img_item = epub.EpubImage(
                uid=f"img_{idx}_{n}", file_name=file_name,
                media_type="image/jpeg", content=img_bytes,
            )
            book.add_item(img_item)
            tags.append(f'<img src="{file_name}" alt="gambar {n}"/>')
        except Exception:
            continue
    return "".join(tags)


def write_combined_epub(path: pathlib.Path, title: str, author: str, story_id: str, results: list,
                         meta: dict = None, cover_bytes: bytes = None, images_by_idx: dict = None,
                         language: str = "id") -> None:
    """Tulis semua chapter yang diberikan ke satu file .epub (satu buku, banyak bab)."""
    meta = meta or {}
    images_by_idx = images_by_idx or {}
    book, epub = _build_epub_book(title, author, story_id, meta=meta, language=language)
    _set_epub_cover(book, epub, cover_bytes)

    intro_body = (
        f"<h1>{html.escape(title)}</h1>"
        f"<p>oleh {html.escape(author)}</p>"
    )
    if meta.get("tags"):
        intro_body += f"<p>Genre: {html.escape(', '.join(meta['tags']))}</p>"
    if meta.get("description"):
        intro_body += f"<p>{html.escape(meta['description'])}</p>"
    intro_body += f"<p>Sumber: https://www.wattpad.com/story/{story_id}</p>"

    intro = epub.EpubHtml(title="Info", file_name="000_info.xhtml", lang=language)
    intro.content = intro_body
    book.add_item(intro)

    epub_chapters = [intro]
    for idx, chapter_title, text in results:
        c = epub.EpubHtml(title=chapter_title, file_name=f"chap_{idx:03d}.xhtml", lang=language)
        img_html = _add_epub_inline_images(book, epub, idx, images_by_idx.get(idx))
        c.content = f"<h1>{html.escape(chapter_title)}</h1>{_text_to_epub_html(text)}{img_html}"
        book.add_item(c)
        epub_chapters.append(c)

    book.toc = tuple(epub_chapters)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav"] + epub_chapters

    epub.write_epub(str(path), book)


def write_separate_epub_zip(path: pathlib.Path, title: str, author: str, story_id: str, results: list,
                             meta: dict = None, cover_bytes: bytes = None, images_by_idx: dict = None,
                             language: str = "id") -> None:
    """Tulis tiap chapter sebagai file .epub terpisah (satu bab per buku), dikemas dalam satu .zip."""
    meta = meta or {}
    images_by_idx = images_by_idx or {}
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = pathlib.Path(tmpdir)

        info_book, epub = _build_epub_book(title, author, story_id, identifier_suffix="info",
                                            meta=meta, language=language)
        _set_epub_cover(info_book, epub, cover_bytes)

        info_body = (
            f"<h1>{html.escape(title)}</h1>"
            f"<p>oleh {html.escape(author)}</p>"
        )
        if meta.get("tags"):
            info_body += f"<p>Genre: {html.escape(', '.join(meta['tags']))}</p>"
        if meta.get("description"):
            info_body += f"<p>{html.escape(meta['description'])}</p>"
        info_body += f"<p>Sumber: https://www.wattpad.com/story/{story_id}</p>"

        info_page = epub.EpubHtml(title="Info", file_name="info.xhtml", lang=language)
        info_page.content = info_body
        info_book.add_item(info_page)
        info_book.toc = (info_page,)
        info_book.add_item(epub.EpubNcx())
        info_book.add_item(epub.EpubNav())
        info_book.spine = ["nav", info_page]
        info_file = tmp_path / "000_info.epub"
        epub.write_epub(str(info_file), info_book)

        chapter_files = [info_file]
        for idx, chapter_title, text in results:
            chap_book, _ = _build_epub_book(
                f"{title} - {chapter_title}", author, story_id, identifier_suffix=f"ch{idx}", language=language
            )
            page = epub.EpubHtml(title=chapter_title, file_name="chapter.xhtml", lang=language)
            img_html = _add_epub_inline_images(chap_book, epub, idx, images_by_idx.get(idx))
            page.content = f"<h1>{html.escape(chapter_title)}</h1>{_text_to_epub_html(text)}{img_html}"
            chap_book.add_item(page)
            chap_book.toc = (page,)
            chap_book.add_item(epub.EpubNcx())
            chap_book.add_item(epub.EpubNav())
            chap_book.spine = ["nav", page]
            fname = tmp_path / f"{idx:03d}_{safe_filename(chapter_title)}.epub"
            epub.write_epub(str(fname), chap_book)
            chapter_files.append(fname)

        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            for f in chapter_files:
                zf.write(f, arcname=f.name)


def _pdf_image_flowable(img_bytes: bytes, max_width: float, max_height: float):
    """Bikin flowable reportlab dari bytes gambar, diskalakan proporsional supaya
    muat dalam max_width x max_height. Kembalikan None kalau gambarnya rusak/tidak
    bisa dibaca — dipanggil dengan asumsi caller melewati None tanpa crash."""
    import io

    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import Image as RLImage

    try:
        reader = ImageReader(io.BytesIO(img_bytes))
        iw, ih = reader.getSize()
        if iw <= 0 or ih <= 0:
            return None
        scale = min(max_width / iw, max_height / ih, 1.0)
        return RLImage(io.BytesIO(img_bytes), width=iw * scale, height=ih * scale)
    except Exception:
        return None


def _text_to_pdf_flowables(text: str, styles) -> list:
    """Ubah teks chapter jadi list flowable Paragraph reportlab, menjaga jeda
    antar paragraf & baris. Karakter khusus di-escape supaya tidak ditafsir
    sebagai mini-markup reportlab."""
    from reportlab.platypus import Paragraph, Spacer

    flowables = []
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        lines = [html.escape(line) for line in para.split("\n")]
        flowables.append(Paragraph("<br/>".join(lines), styles["BodyText"]))
        flowables.append(Spacer(1, 8))
    return flowables


def _build_pdf_story_intro(title, author, story_id, meta, styles) -> list:
    """Bagian halaman info/intro (judul, penulis, genre, sinopsis) yang dipakai
    baik di PDF gabungan maupun di file info.pdf pada mode terpisah/zip."""
    from reportlab.platypus import Paragraph, Spacer

    meta = meta or {}
    flowables = [
        Paragraph(html.escape(title), styles["Title"]),
        Paragraph(f"oleh {html.escape(author)}", styles["Normal"]),
        Spacer(1, 8),
    ]
    if meta.get("tags"):
        flowables.append(Paragraph(f"Genre: {html.escape(', '.join(meta['tags']))}", styles["Normal"]))
    if meta.get("description"):
        flowables.append(Paragraph(html.escape(meta["description"]), styles["Normal"]))
    flowables.append(Paragraph(f"Sumber: https://www.wattpad.com/story/{story_id}", styles["Normal"]))
    return flowables


def write_combined_pdf(path: pathlib.Path, title: str, author: str, story_id: str, results: list,
                        meta: dict = None, cover_bytes: bytes = None, images_by_idx: dict = None) -> None:
    """Tulis semua chapter yang diberikan ke satu file .pdf."""
    from reportlab.lib.pagesizes import A5
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import PageBreak, SimpleDocTemplate

    images_by_idx = images_by_idx or {}
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=A5, title=title, author=author)
    story = []

    if cover_bytes:
        cover_flowable = _pdf_image_flowable(cover_bytes, doc.width, doc.height * 0.85)
        if cover_flowable:
            story.append(cover_flowable)
            story.append(PageBreak())

    story.extend(_build_pdf_story_intro(title, author, story_id, meta, styles))
    story.append(PageBreak())

    for idx, chapter_title, text in results:
        story.append(_pdf_chapter_heading(chapter_title, styles))
        story.extend(_text_to_pdf_flowables(text, styles))
        for img_bytes in images_by_idx.get(idx) or []:
            flowable = _pdf_image_flowable(img_bytes, doc.width, doc.height * 0.7)
            if flowable:
                story.append(flowable)
        story.append(PageBreak())

    doc.build(story)


def _pdf_chapter_heading(text: str, styles):
    """Kecil: heading chapter, dipisah jadi fungsi supaya gampang dipakai ulang
    di writer PDF gabungan maupun terpisah tanpa impor berulang."""
    from reportlab.platypus import Paragraph

    return Paragraph(html.escape(text), styles["Heading1"])


def write_separate_pdf_zip(path: pathlib.Path, title: str, author: str, story_id: str, results: list,
                            meta: dict = None, cover_bytes: bytes = None, images_by_idx: dict = None) -> None:
    """Tulis tiap chapter sebagai file .pdf terpisah, dikemas dalam satu .zip."""
    from reportlab.lib.pagesizes import A5
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate

    images_by_idx = images_by_idx or {}
    styles = getSampleStyleSheet()
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = pathlib.Path(tmpdir)

        info_file = tmp_path / "000_info.pdf"
        info_doc = SimpleDocTemplate(str(info_file), pagesize=A5, title=title, author=author)
        info_story = []
        if cover_bytes:
            cover_flowable = _pdf_image_flowable(cover_bytes, info_doc.width, info_doc.height * 0.85)
            if cover_flowable:
                info_story.append(cover_flowable)
        info_story.extend(_build_pdf_story_intro(title, author, story_id, meta, styles))
        info_doc.build(info_story)

        chapter_files = [info_file]
        for idx, chapter_title, text in results:
            fname = tmp_path / f"{idx:03d}_{safe_filename(chapter_title)}.pdf"
            chap_doc = SimpleDocTemplate(str(fname), pagesize=A5, title=chapter_title, author=author)
            chap_story = [_pdf_chapter_heading(chapter_title, styles)]
            chap_story.extend(_text_to_pdf_flowables(text, styles))
            for img_bytes in images_by_idx.get(idx) or []:
                flowable = _pdf_image_flowable(img_bytes, chap_doc.width, chap_doc.height * 0.7)
                if flowable:
                    chap_story.append(flowable)
            chap_doc.build(chap_story)
            chapter_files.append(fname)

        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            for f in chapter_files:
                zf.write(f, arcname=f.name)


def check_pdf_available(console=None) -> None:
    """Pastikan library reportlab terinstall sebelum dipakai."""
    try:
        import reportlab  # noqa: F401
    except ImportError:
        if console:
            console.print("\n[danger]❌  Library 'reportlab' belum terinstall.[/danger]")
            console.print("    [muted]Jalankan:[/muted] [accent]pip install reportlab[/accent]")
        else:
            print("Library 'reportlab' belum terinstall. Jalankan: pip install reportlab")
        sys.exit(1)
